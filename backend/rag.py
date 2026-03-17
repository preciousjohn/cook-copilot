"""
RAG (Retrieval-Augmented Generation) module for CookCopilot.

Architecture:
  1. Load   — read KB files from disk (.md, .txt, .pdf, .docx)
  2. Chunk  — split into semantically meaningful sections
  3. Embed  — compute OpenAI embeddings for each chunk
  4. Store  — persist in ChromaDB (local, file-based)
  5. Query  — retrieve top-k relevant chunks for a given query
  6. Format — return retrieved context as a string for LLM injection

Each knowledge base (dietitian, chef, engineer, …) gets its own
ChromaDB *collection*, so they can be built / queried independently.

Supported file types:
  - .md / .txt  — read as UTF-8 text, split by ## headings
  - .pdf        — text extracted via pymupdf (pip install pymupdf)
  - .docx       — text extracted via python-docx (pip install python-docx)
"""

from __future__ import annotations

import hashlib
import os
import re
import shutil                          # ← FIX 1: added for rmtree
from pathlib import Path
from typing import List, Dict, Optional

from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from dotenv import load_dotenv
from openai import OpenAI

load_dotenv()
client = OpenAI()

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

BACKEND_DIR = Path(__file__).resolve().parent
CHROMA_PERSIST_DIR = BACKEND_DIR / "vectorstores"

# Ensure persist directory exists
CHROMA_PERSIST_DIR.mkdir(parents=True, exist_ok=True)



# ---------------------------------------------------------------------------
# 1. Load KB text  (supports .md, .txt, .pdf, .docx)
# ---------------------------------------------------------------------------

def _extract_pdf_text(filepath: Path) -> str:
    """Extract text from a PDF using pymupdf (fitz)."""
    try:
        import fitz  # pymupdf
    except ImportError:
        raise ImportError(
            "pymupdf is required for PDF support. "
            "Install it with: pip install pymupdf"
        )
    doc = fitz.open(str(filepath))
    pages: List[str] = []
    for page in doc:
        text = page.get_text("text")
        if text.strip():
            pages.append(text)
    doc.close()
    return "\n\n".join(pages)


def _extract_docx_text(filepath: Path) -> str:
    """Extract text from a DOCX using python-docx."""
    try:
        import docx
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX support. "
            "Install it with: pip install python-docx"
        )
    doc = docx.Document(str(filepath))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


def load_kb_file(path: str) -> str:
    """
    Read a knowledge-base file. Resolves relative paths to backend/.
    Supported formats: .md, .txt, .pdf, .docx
    """
    p = Path(path)
    if not p.is_absolute():
        p = BACKEND_DIR / p
    if not p.exists():
        return ""

    ext = p.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf_text(p)
    elif ext == ".docx":
        return _extract_docx_text(p)
    else:
        # .md, .txt, or any other text file
        return p.read_text(encoding="utf-8")


# ---------------------------------------------------------------------------
# 2. Chunk — heading-aware splitter (works for MD and plain text)
# ---------------------------------------------------------------------------

def _heading_aware_split(text: str, chunk_size: int = 500, chunk_overlap: int = 50) -> List[Document]:
    """
    Split text into chunks for embedding.

    Strategy:
    - If the text contains ## headings (markdown), split by heading first,
      then sub-split large sections with RecursiveCharacterTextSplitter.
    - If no headings found (e.g. PDF-extracted text), fall back to
      paragraph-based recursive splitting directly.
    """
    has_headings = bool(re.search(r"^## ", text, flags=re.MULTILINE))

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    docs: List[Document] = []

    if has_headings:
        # Markdown path: split by ## headings
        sections = re.split(r"(?=^## )", text, flags=re.MULTILINE)
        sections = [s.strip() for s in sections if s.strip()]

        for section in sections:
            lines = section.split("\n", 1)
            heading = lines[0].strip().lstrip("#").strip() if lines else "unknown"

            if len(section) <= chunk_size:
                docs.append(Document(
                    page_content=section,
                    metadata={"heading": heading},
                ))
            else:
                sub_docs = splitter.create_documents(
                    [section],
                    metadatas=[{"heading": heading}],
                )
                docs.extend(sub_docs)
    else:
        # Plain-text path (PDF, DOCX, etc.): paragraph-based splitting
        sub_docs = splitter.create_documents(
            [text],
            metadatas=[{"heading": "document"}],
        )
        docs.extend(sub_docs)

    return docs


# ---------------------------------------------------------------------------
# 3–4. Embed + Store  (build or load ChromaDB collection)
# ---------------------------------------------------------------------------

def _collection_name(kb_path: str) -> str:
    """Deterministic collection name from the KB file path.
    ChromaDB requires 3-512 chars from [a-zA-Z0-9._-], starting/ending alphanumeric.
    """
    stem = Path(kb_path).stem  # e.g. "Lecture 10 - Estimating Nutrient Needs"
    # Replace spaces and invalid chars with underscores
    sanitized = re.sub(r"[^a-zA-Z0-9._-]", "_", stem)
    # Collapse multiple underscores
    sanitized = re.sub(r"_+", "_", sanitized).strip("_")
    # Ensure starts/ends with alphanumeric
    sanitized = sanitized.strip("._-") or "kb"
    name = f"kb_{sanitized}"
    # Truncate to 512 chars
    return name[:512]


def _content_hash(text: str) -> str:
    """SHA-256 hash of the KB content — used to detect changes."""
    return hashlib.sha256(text.encode("utf-8")).hexdigest()[:16]


class KnowledgeBaseStore:
    """
    Manages vector stores for one or more knowledge bases.
    Caches loaded collections in memory; persists embeddings to disk.
    """

    def __init__(self):
        self._embeddings = OpenAIEmbeddings(
            model="text-embedding-3-small"
        )
        self._stores: Dict[str, Chroma] = {}
        self._hashes: Dict[str, str] = {}

    # ---- public API -------------------------------------------------------

    def index_kb(self, kb_path: str, *, force: bool = False) -> int:
        """
        Build (or rebuild) the vector index for a KB file.
        Returns the number of chunks indexed.
        Skips if the content hash hasn't changed (unless force=True).
        """
        text = load_kb_file(kb_path)
        if not text:
            raise FileNotFoundError(f"KB not found: {kb_path}")

        col_name = _collection_name(kb_path)
        new_hash = _content_hash(text)

        # Skip if unchanged
        if not force and col_name in self._hashes and self._hashes[col_name] == new_hash:
            return 0

        # Chunk
        docs = _heading_aware_split(text)

        # Add source metadata
        for doc in docs:
            doc.metadata["source"] = kb_path

        # Build ChromaDB collection (overwrites if exists)
        persist_dir = str(CHROMA_PERSIST_DIR / col_name)

        # FIX 1: Delete existing collection directory before re-indexing
        # to prevent duplicate chunks accumulating across server restarts.
        # (_hashes is in-memory and cleared on restart, but ChromaDB persists
        #  on disk, so from_documents would ADD to existing chunks.)
        if Path(persist_dir).exists():
            shutil.rmtree(persist_dir)

        store = Chroma.from_documents(
            documents=docs,
            embedding=self._embeddings,
            collection_name=col_name,
            persist_directory=persist_dir,
        )

        self._stores[col_name] = store
        self._hashes[col_name] = new_hash

        return len(docs)

    def retrieve(
        self,
        kb_path: str,
        query: str,
        top_k: int = 5,
    ) -> List[Document]:
        """
        Retrieve the top-k most relevant chunks from a KB's vector store.
        Automatically indexes the KB if not yet loaded.
        """
        col_name = _collection_name(kb_path)

        # Auto-index if not loaded yet
        if col_name not in self._stores:
            self.index_kb(kb_path)

        store = self._stores[col_name]
        results = store.similarity_search(query, k=top_k)
        return results

    def retrieve_with_scores(
        self,
        kb_path: str,
        query: str,
        top_k: int = 5,
    ) -> List[Dict]:
        """
        Like retrieve(), but also returns similarity scores.
        Returns list of {"content": str, "metadata": dict, "score": float}.
        """
        col_name = _collection_name(kb_path)

        if col_name not in self._stores:
            self.index_kb(kb_path)

        store = self._stores[col_name]
        # Use raw distance scores (cosine distance: 0=identical, 2=opposite)
        # and convert to similarity (1 - distance) to get values in [0, 1].
        results = store.similarity_search_with_score(query, k=top_k)

        return [
            {
                "content": doc.page_content,
                "metadata": doc.metadata,
                "score": round(1.0 - distance, 4),
            }
            for doc, distance in results
        ]

    def format_context(
        self,
        kb_path: str,
        query: str,
        top_k: int = 5,
        min_score: float = 0.2,
    ) -> str:
        """
        Retrieve chunks and format them into a single context string
        ready for LLM system-prompt injection.
        Only includes chunks with cosine similarity >= min_score.
        """
        chunks = self.retrieve_with_scores(kb_path, query, top_k=top_k)
        relevant = [c for c in chunks if c["score"] >= min_score]
        if not relevant:
            return ""

        parts: List[str] = []
        for i, chunk in enumerate(relevant, 1):
            heading = chunk["metadata"].get("heading", "")
            parts.append(
                f"[Retrieved Chunk {i} | Section: {heading} | score={chunk['score']:.2f}]\n"
                f"{chunk['content']}"
            )
        return "\n\n---\n\n".join(parts)

    def list_all_chunks(self, kb_path: str) -> List[Dict]:
        """
        Return ALL chunks stored in ChromaDB for a given KB.
        Useful for debugging / visualizing the vector store contents.
        """
        col_name = _collection_name(kb_path)

        if col_name not in self._stores:
            self.index_kb(kb_path)

        store = self._stores[col_name]
        collection = store._collection
        result = collection.get(include=["documents", "metadatas", "embeddings"])

        chunks = []
        ids = result.get("ids", [])
        docs = result.get("documents", [])
        metas = result.get("metadatas", [])
        embeddings = result.get("embeddings", [])

        for i in range(len(ids)):
            # FIX 2: embeddings can be a numpy array — bare `if embeddings`
            # raises "truth value of array is ambiguous". Check explicitly.
            has_emb = (
                embeddings is not None
                and len(embeddings) > i
                and embeddings[i] is not None
            )
            chunk = {
                "id": ids[i],
                "content": docs[i] if docs else None,
                "metadata": metas[i] if metas else None,
                "embedding_dim": len(embeddings[i]) if has_emb else None,
                "embedding_preview": list(embeddings[i][:5]) if has_emb else None,
            }
            chunks.append(chunk)

        return chunks


# ---------------------------------------------------------------------------
# Module-level singleton
# ---------------------------------------------------------------------------

kb_store = KnowledgeBaseStore()